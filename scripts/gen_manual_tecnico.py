# -*- coding: utf-8 -*-
"""
Genera la plantilla del Manual Tecnico (estructura + formato) en formato .docx
para el proyecto Marketplace-Residencia.
Salida: docs/Manual-Tecnico-Marketplace-Residencia.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Paleta ----------
AZUL = RGBColor(0x1F, 0x3A, 0x5F)
AZUL_CLARO = RGBColor(0x2E, 0x5A, 0x88)
GRIS = RGBColor(0x55, 0x55, 0x55)
VERDE = RGBColor(0x1E, 0x7D, 0x32)
AMBAR = RGBColor(0xB7, 0x6E, 0x00)
ROJO = RGBColor(0xB0, 0x2A, 0x2A)

doc = Document()

# ---------- Estilos base ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
normal.paragraph_format.line_spacing = 1.15

for hname, size, color in [("Heading 1", 16, AZUL), ("Heading 2", 13, AZUL_CLARO),
                           ("Heading 3", 11.5, GRIS)]:
    st = doc.styles[hname]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True

# ---------- Helpers ----------

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=10, white=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color:
        run.font.color.rgb = color

def add_table(headers, rows, widths=None, header_fill="1F3A5F"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, white=True, size=10)
        shade(hdr[i], header_fill)
    for r in rows:
        cells = t.add_row().cells
        for i, val in enumerate(r):
            set_cell_text(cells[i], str(val), size=9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def guidance_box(objetivo, contenido, extension, diagramas=None, tablas=None, capturas=None):
    """Caja-guia sombreada que indica que debe contener el capitulo."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    shade(cell, "EAF1F8")
    cell.text = ""

    def line(label, value, lcolor=AZUL):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(label + ": ")
        r1.font.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = lcolor
        r2 = p.add_run(value)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = GRIS

    # primer parrafo reemplaza el vacio inicial
    cell.paragraphs[0].text = ""
    p0 = cell.paragraphs[0]
    r = p0.add_run("GUIA DE REDACCION")
    r.font.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = AZUL_CLARO

    line("Objetivo", objetivo)
    line("Contenido a incluir", contenido)
    line("Extension recomendada", extension, lcolor=VERDE)
    if diagramas:
        line("Diagramas", diagramas, lcolor=AMBAR)
    if tablas:
        line("Tablas", tablas, lcolor=AMBAR)
    if capturas:
        line("Capturas", capturas, lcolor=AMBAR)
    doc.add_paragraph()

def placeholder(text="[ Contenido por redactar ]"):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    r.font.size = Pt(10.5)

def chapter(num, title):
    doc.add_page_break()
    h = doc.add_heading(level=1)
    run = h.add_run(f"{num}. {title}")
    run.font.bold = True

def add_field(paragraph, field_code):
    """Inserta un campo de Word (TOC, PAGE, etc.)."""
    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = field_code
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    fldText = OxmlElement("w:t"); fldText.text = "Actualice este campo (clic derecho > Actualizar campos)"
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin); run._r.append(instr); run._r.append(fldSep)
    run._r.append(fldText); run._r.append(fldEnd)

# ---------- Pie de pagina con numero ----------
section = doc.sections[0]
section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
footer_p = section.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run("Manual Tecnico - Marketplace-Residencia   |   Pagina ")
fr.font.size = Pt(8); fr.font.color.rgb = GRIS
add_field(footer_p, "PAGE")

# ======================================================================
# PORTADA
# ======================================================================
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MANUAL TECNICO"); r.font.size = Pt(34); r.font.bold = True; r.font.color.rgb = AZUL
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Marketplace-Residencia"); r.font.size = Pt(20); r.font.color.rgb = AZUL_CLARO
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Plataforma de comercio electronico multi-vendedor (mezcal de Oaxaca)")
r.font.size = Pt(12); r.font.italic = True; r.font.color.rgb = GRIS
for _ in range(6):
    doc.add_paragraph()

meta = doc.add_table(rows=6, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
labels = ["Autor(es)", "Institucion / Empresa", "Asesor(es)", "Version del documento", "Fecha", "Confidencialidad"]
vals = ["[Nombre]", "[Institucion]", "[Asesor academico / lider tecnico]", "1.0", "[DD/MM/AAAA]", "Uso interno"]
for i, (l, v) in enumerate(zip(labels, vals)):
    set_cell_text(meta.rows[i].cells[0], l, bold=True, color=AZUL, size=11)
    set_cell_text(meta.rows[i].cells[1], v, size=11)
for row in meta.rows:
    row.cells[0].width = Cm(5); row.cells[1].width = Cm(9)

# ======================================================================
# CONTROL DE VERSIONES
# ======================================================================
doc.add_page_break()
doc.add_heading("Control de versiones", level=1)
add_table(
    ["Version", "Fecha", "Autor", "Descripcion del cambio"],
    [["1.0", "[DD/MM/AAAA]", "[Nombre]", "Version inicial del manual"],
     ["", "", "", ""],
     ["", "", "", ""]],
    widths=[2.2, 3, 3.5, 6])

# ======================================================================
# INDICE
# ======================================================================
doc.add_heading("Indice de contenido", level=1)
note = doc.add_paragraph()
nr = note.add_run("Tabla de contenido automatica. En Word: clic derecho sobre el indice > "
                  "\"Actualizar campos\" > \"Actualizar toda la tabla\".")
nr.font.italic = True; nr.font.size = Pt(9); nr.font.color.rgb = GRIS
toc_p = doc.add_paragraph()
add_field(toc_p, 'TOC \\o "1-2" \\h \\z \\u')

# Resumen de distribucion de paginas
doc.add_page_break()
doc.add_heading("Distribucion de paginas (objetivo: 50)", level=1)
add_table(
    ["#", "Capitulo", "Paginas"],
    [["1", "Introduccion y alcance", "2"],
     ["2", "Arquitectura general", "4"],
     ["3", "Stack tecnologico y dependencias", "2"],
     ["4", "Modelo de datos", "5"],
     ["5", "Modulos del backend", "6"],
     ["6", "Estructura del frontend", "4"],
     ["7", "Roles, permisos y seguridad", "5"],
     ["8", "Flujos de negocio criticos", "6"],
     ["9", "Integraciones externas", "4"],
     ["10", "Internacionalizacion, moneda e impuestos", "2"],
     ["11", "Configuracion, variables de entorno y despliegue", "4"],
     ["12", "Operacion y mantenimiento", "2"],
     ["13", "Limitaciones, deuda tecnica y trabajo futuro", "2"],
     ["14", "Anexos", "2"],
     ["", "TOTAL", "50"]],
    widths=[1.5, 11, 2.5])

# ======================================================================
# CAPITULO 1
# ======================================================================
chapter(1, "Introduccion y alcance")
guidance_box(
    objetivo="Explicar que es el sistema y que cubre (y que no) este manual.",
    contenido="Proposito del sistema, alcance funcional, publico objetivo del manual, "
              "convenciones tipograficas y glosario breve de terminos del dominio.",
    extension="2 paginas",
    tablas="Estado de funcionalidades (incluida abajo, ya verificada contra el codigo).")
doc.add_heading("1.1 Proposito del sistema", level=2); placeholder()
doc.add_heading("1.2 Alcance del manual", level=2); placeholder()
doc.add_heading("1.3 Estado de funcionalidades", level=2)
add_table(
    ["Funcionalidad", "Estado", "Ubicacion / Evidencia"],
    [["Marketplace multi-vendedor", "Implementado", "detalle_pedido + pedido_productor"],
     ["Gestion de usuarios", "Implementado", "modulo usuarios (CRUD + soft delete)"],
     ["Gestion de productores", "Implementado", "modulo productores (onboarding)"],
     ["Catalogo de productos", "Implementado", "modulo productos + categorias"],
     ["Inventario", "Implementado", "modulo inventario + movimientos + lotes"],
     ["Carrito de compras", "Implementado", "modulo carrito + CarritoContext"],
     ["Checkout", "Implementado", "tienda/checkout (4 pasos)"],
     ["Pasarela de pagos", "Implementado", "Stripe + PayPal (reales)"],
     ["Gestion de ordenes", "Implementado", "modulo pedidos (ciclo completo)"],
     ["Direcciones", "Implementado", "modulo direcciones (MX + US)"],
     ["Envios nacionales", "Implementado", "SkydropX nacional"],
     ["Envios internacionales", "Implementado", "SkydropX intl + aduana (HS)"],
     ["Integracion SkydropX", "Implementado", "skydropx.service.ts"],
     ["Gestion de impuestos", "Implementado", "tasas_impuesto (IVA/IEPS + US)"],
     ["Conversion de monedas", "Implementado", "tasas-cambio + cron FX"],
     ["i18n (frontend)", "Implementado", "next-intl es/en"],
     ["i18n (backend)", "Parcial", "modulo i18n stub"],
     ["Panel administrativo", "Implementado", "administrador/* (~17 paginas)"],
     ["Dashboard", "Implementado", "KPIs admin + productor"],
     ["Reportes", "Parcial", "graficas si; export consolidado no"],
     ["Notificaciones", "Implementado", "modulo notificaciones + email"],
     ["Seguridad", "Implementado", "JWT, RBAC, throttling, CORS, XSS"]],
    widths=[5.5, 2.5, 7])

# ======================================================================
# CAPITULO 2
# ======================================================================
chapter(2, "Arquitectura general")
guidance_box(
    objetivo="Dar una vision de alto nivel de como esta organizado y como fluye una peticion.",
    contenido="Monorepo Turborepo (apps/api, apps/web, packages/database), responsabilidades "
              "de cada parte, comunicacion frontend-backend (proxy de Next a la API), patron general.",
    extension="4 paginas",
    diagramas="Diagrama de componentes, diagrama de despliegue, diagrama de flujo de una peticion.",
    capturas="(Opcional) estructura de carpetas del repositorio.")
doc.add_heading("2.1 Vision de componentes", level=2); placeholder("[ Insertar diagrama de arquitectura ]")
doc.add_heading("2.2 Estructura del monorepo", level=2); placeholder()
doc.add_heading("2.3 Flujo de una peticion (proxy Next -> API)", level=2); placeholder("[ Insertar diagrama de flujo ]")

# ======================================================================
# CAPITULO 3
# ======================================================================
chapter(3, "Stack tecnologico y dependencias")
guidance_box(
    objetivo="Listar el stack y las dependencias notables con su justificacion breve.",
    contenido="Versiones de NestJS 11, Next 15, React 19, Prisma 5.22, Node; dependencias "
              "agrupadas por concern (auth, pagos, envios, UI, datos).",
    extension="2 paginas",
    tablas="Tabla de dependencias notables por categoria (plantilla incluida).")
add_table(
    ["Categoria", "Libreria", "Version", "Uso"],
    [["Backend core", "@nestjs/* ", "^11", "Framework API"],
     ["ORM / BD", "@prisma/client", "^5.22", "Acceso a Postgres (Neon)"],
     ["Pagos", "stripe / @paypal/paypal-server-sdk", "^22 / ^2.3", "Cobros y payouts"],
     ["Frontend core", "next / react", "^15.3 / ^19.2", "App Router SSR/CSR"],
     ["i18n", "next-intl", "^4.9", "Locales es/en"],
     ["UI pagos", "@stripe/react-stripe-js / @paypal/react-paypal-js", "^6.3 / ^9.2", "Checkout"],
     ["Graficas", "recharts", "^3.8", "Dashboards"],
     ["[ ... ]", "[ ... ]", "[ ... ]", "[ ... ]"]],
    widths=[3, 5.5, 3, 4])

# ======================================================================
# CAPITULO 4
# ======================================================================
chapter(4, "Modelo de datos")
guidance_box(
    objetivo="Documentar el esquema de base de datos para mantenimiento y evolucion.",
    contenido="Proveedor (PostgreSQL/Neon, pooler vs directo), modelos agrupados por dominio "
              "(~55 tablas), tipos especiales (BigInt, Decimal, JSON/GIN, cifrado), soft deletes, "
              "estrategia de migraciones.",
    extension="5 paginas",
    diagramas="Diagrama ER dividido por dominios (auth, catalogo/inventario, ordenes, envios, finanzas).",
    tablas="Inventario de modelos por dominio (plantilla incluida).")
doc.add_heading("4.1 Diagrama Entidad-Relacion", level=2)
placeholder("[ Insertar ER generado desde packages/database/prisma/schema.prisma ]")
doc.add_heading("4.2 Modelos por dominio", level=2)
add_table(
    ["Dominio", "Modelos principales"],
    [["Auth y usuarios", "usuarios, usuario_rol, roles, oauth_cuentas, refresh_tokens, auditoria"],
     ["Catalogo e inventario", "productos, producto_imagenes, categorias, inventario, movimientos_inventario, lotes"],
     ["Tiendas y productores", "tiendas, productores, productor_categoria, regiones"],
     ["Ordenes", "pedidos, detalle_pedido, pedido_productor, pagos, facturas"],
     ["Envios", "envios, envio_guias, envio_eventos, servicios_envio, transportistas, integraciones_envio"],
     ["Finanzas", "comisiones, tasas_cambio, tasas_impuesto, payouts, payment_fees, refunds"],
     ["Sistema y global", "configuracion_sistema, paises, idiomas, direcciones, notificaciones, resenas, webhook_events_log"]],
    widths=[4, 11])
doc.add_heading("4.3 Tipos especiales y convenciones", level=2)
placeholder("[ BigInt como ids, Decimal para montos/tasas, JSON/GIN, Bytes cifrados, soft delete eliminado_en ]")
doc.add_heading("4.4 Estrategia de migraciones", level=2); placeholder()

# ======================================================================
# CAPITULO 5
# ======================================================================
chapter(5, "Modulos del backend")
guidance_box(
    objetivo="Mapear las responsabilidades de cada modulo del backend NestJS.",
    contenido="Tabla de los 35 modulos (responsabilidad, endpoints principales, estado), patron "
              "module/controller/service/dto y concerns transversales (ValidationPipe, serializacion "
              "BigInt, filtro de excepciones, interceptor de sanitizacion).",
    extension="6 paginas",
    tablas="Inventario de modulos (plantilla con encabezado; completar con los 35).")
add_table(
    ["Modulo", "Responsabilidad", "Endpoints clave", "Estado"],
    [["auth", "Login/registro/refresh, OAuth, reset password", "POST /auth/login, /auth/refresh", "Impl."],
     ["usuarios", "Perfiles y CRUD de usuarios", "CRUD /usuarios", "Impl."],
     ["pagos", "Stripe + PayPal, webhooks, disputas", "/pagos/stripe/intent, /pagos/*/webhook", "Impl."],
     ["pedidos", "Ciclo de vida de la orden", "CRUD /pedidos, transiciones estado", "Impl."],
     ["envios", "SkydropX cotizacion/guia/tracking", "/envios/cotizar, /envios/guia", "Impl."],
     ["payouts", "Dispersion a productores", "/payouts/generar", "Impl."],
     ["i18n", "Traducciones backend", "-", "Stub"],
     ["[ ... completar resto de modulos ... ]", "", "", ""]],
    widths=[2.8, 5, 4.7, 2])
doc.add_heading("5.1 Patron de modulo (module/controller/service/dto)", level=2); placeholder()
doc.add_heading("5.2 Concerns transversales", level=2)
placeholder("[ ValidationPipe global, BigInt.toJSON, AllExceptionsFilter, SanitizeBodyInterceptor ]")

# ======================================================================
# CAPITULO 6
# ======================================================================
chapter(6, "Estructura del frontend")
guidance_box(
    objetivo="Explicar la organizacion del cliente web Next.js y como se conecta a la API.",
    contenido="Rutas agrupadas por rol (publico, cliente, productor, admin), seleccion de layout, "
              "proveedores de contexto, cliente API (api.ts) con refresh singleton, rewrites de proxy.",
    extension="4 paginas",
    tablas="Mapa de rutas por rol (plantilla incluida).",
    capturas="Dashboard admin, panel de productor (1-2 capturas representativas).")
add_table(
    ["Rol / Grupo", "Rutas principales"],
    [["Publico", "auth/*, cliente/producto, categoria/[slug], paginas legales"],
     ["Cliente", "cliente/pedidos, direcciones, tienda/carrito, tienda/checkout, tienda/deseos"],
     ["Productor", "dashboard/productor/* (productos, pedidos, lotes, ventas, ingresos, tienda)"],
     ["Admin", "administrador/* (usuarios, productores, pedidos, payouts, comisiones, roles-permisos)"]],
    widths=[3.5, 11.5])
doc.add_heading("6.1 Seleccion de layout", level=2); placeholder()
doc.add_heading("6.2 Proveedores de contexto", level=2)
placeholder("[ Auth, Carrito, Wishlist, Config, Locale, Sidebar, Theme, Session ]")
doc.add_heading("6.3 Cliente API y refresco de token", level=2); placeholder()

# ======================================================================
# CAPITULO 7
# ======================================================================
chapter(7, "Roles, permisos y seguridad")
guidance_box(
    objetivo="Documentar el modelo de seguridad y control de acceso.",
    contenido="RBAC (roles/permisos en BD, payload del JWT), guards (Auth/Roles/Permisos), JWT manual "
              "HMAC-SHA256, throttling, CORS, headers de seguridad, sanitizacion XSS, manejo de secretos/cifrado.",
    extension="5 paginas",
    diagramas="Diagrama de como Auth/Roles/Permisos filtran una peticion.",
    tablas="Matriz roles x permisos (plantilla incluida).")
doc.add_heading("7.1 Modelo RBAC", level=2)
add_table(
    ["Rol", "Descripcion", "Permisos representativos"],
    [["admin", "Operador del marketplace", "Gestion global, payouts, configuracion"],
     ["productor", "Vendedor (productor de mezcal)", "Sus productos, pedidos, lotes, ingresos"],
     ["cliente", "Comprador final", "Compras, carrito, direcciones, resenas"],
     ["[ ... ]", "", ""]],
    widths=[2.5, 5, 7.5])
doc.add_heading("7.2 Guards y flujo de autorizacion", level=2); placeholder()
doc.add_heading("7.3 JWT, throttling, CORS y headers", level=2); placeholder()
doc.add_heading("7.4 Manejo de secretos y cifrado", level=2); placeholder()

# ======================================================================
# CAPITULO 8
# ======================================================================
chapter(8, "Flujos de negocio criticos")
guidance_box(
    objetivo="Documentar como opera el dinero y la mercancia de extremo a extremo.",
    contenido="Diagramas de secuencia de: (a) carrito->checkout->pago, (b) webhook->orden pagada->"
              "pedido_productor, (c) payouts admin->transfers, (d) cotizacion->guia->tracking; "
              "maquina de estados del pedido.",
    extension="6 paginas",
    diagramas="4 diagramas de secuencia + 1 maquina de estados del pedido.",
    capturas="Pasos del checkout, detalle/tracking de pedido, vista de payouts.")
doc.add_heading("8.1 Checkout y pago", level=2); placeholder("[ Insertar diagrama de secuencia ]")
doc.add_heading("8.2 Webhook -> orden pagada -> dispersion por productor", level=2); placeholder("[ Diagrama ]")
doc.add_heading("8.3 Generacion de payouts", level=2); placeholder("[ Diagrama ]")
doc.add_heading("8.4 Cotizacion, guia y tracking de envio", level=2); placeholder("[ Diagrama ]")
doc.add_heading("8.5 Maquina de estados del pedido", level=2)
placeholder("[ pendiente -> pagado -> procesando -> enviado -> entregado | expirado | cancelado ]")

# ======================================================================
# CAPITULO 9
# ======================================================================
chapter(9, "Integraciones externas")
guidance_box(
    objetivo="Documentar las dependencias de terceros y su operacion/soporte.",
    contenido="Stripe (Connect, tax, disputas, refunds), PayPal (orders/payouts), SkydropX (cotizacion, "
              "guia, aduana HS, restricciones de alcohol), SendGrid/Gmail fallback, ExchangeRate-API, "
              "Google OAuth. Credenciales, modos sandbox/produccion, manejo de errores y webhooks.",
    extension="4 paginas",
    tablas="Matriz de integraciones (plantilla incluida).")
add_table(
    ["Servicio", "Uso", "Credenciales (env)", "Modo", "Notas"],
    [["Stripe", "Cobros, Connect, payouts, disputas", "STRIPE_SECRET_KEY, STRIPE_WEBHOOK_SECRET", "test/live", "Webhooks dedup"],
     ["PayPal", "Orders, capture, payouts", "PAYPAL_CLIENT_ID/SECRET", "sandbox/live", "Verif. firma webhook"],
     ["SkydropX", "Cotizacion, guia, aduana", "SKYDROPX_CLIENT_ID/SECRET", "sandbox/prod", "OAuth2 + polling"],
     ["SendGrid", "Email transaccional", "SENDGRID_API_KEY, EMAIL_FROM", "-", "Fallback Gmail SMTP"],
     ["ExchangeRate-API", "Sync de tasas FX", "EXCHANGERATE_API_KEY", "-", "Cron horario"],
     ["Google OAuth", "Login social", "GOOGLE_CLIENT_ID/SECRET", "-", "Via NextAuth"]],
    widths=[2.5, 3.8, 4.5, 1.8, 2.4])

# ======================================================================
# CAPITULO 10
# ======================================================================
chapter(10, "Internacionalizacion, moneda e impuestos")
guidance_box(
    objetivo="Documentar las reglas multinacionales (MX/US) que no son obvias en el codigo.",
    contenido="Locales es/en, mapeo locale->moneda, conversion MXN<->USD, IVA/IEPS incluido en precio, "
              "impuestos US, restricciones de alcohol por estado.",
    extension="2 paginas",
    tablas="Reglas de moneda por locale/pais; restricciones de alcohol por estado.")
doc.add_heading("10.1 Locales y moneda", level=2); placeholder()
doc.add_heading("10.2 Impuestos (MX: IVA/IEPS; US)", level=2); placeholder()
doc.add_heading("10.3 Restricciones de alcohol por estado", level=2); placeholder()

# ======================================================================
# CAPITULO 11
# ======================================================================
chapter(11, "Configuracion, variables de entorno y despliegue")
guidance_box(
    objetivo="Permitir levantar el entorno local y desplegar a produccion.",
    contenido="Inventario de variables de entorno por dominio, archivos .env por app, scripts de Turbo, "
              "build/start, despliegue (Vercel web / host Node para la API), assets estaticos /uploads, "
              "pooler Neon vs conexion directa.",
    extension="4 paginas",
    tablas="Inventario de variables de entorno (plantilla incluida).")
doc.add_heading("11.1 Variables de entorno", level=2)
add_table(
    ["Dominio", "Variables"],
    [["Base de datos", "DATABASE_URL, DIRECT_URL"],
     ["Auth / JWT", "JWT_ACCESS_SECRET, JWT_REFRESH_SECRET, NEXTAUTH_SECRET, NEXTAUTH_URL"],
     ["Google OAuth", "GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET, GOOGLE_REDIRECT_URI"],
     ["Pagos", "STRIPE_SECRET_KEY, STRIPE_WEBHOOK_SECRET, STRIPE_CONNECT_CLIENT_ID, PAYPAL_CLIENT_ID/SECRET, PAYPAL_MODE"],
     ["Envios", "SKYDROPX_ENV, SKYDROPX_CLIENT_ID/SECRET, SKYDROPX_SHIPPER_*, SKYDROPX_HS_CODE_DEFAULT"],
     ["Email", "SENDGRID_API_KEY, EMAIL_FROM, GMAIL_USER, GMAIL_APP_PASSWORD"],
     ["Moneda/FX", "EXCHANGERATE_API_KEY, DEFAULT_CURRENCY_USA, FX_MAX_AGE_HORAS"],
     ["URLs / CORS", "FRONTEND_URL, NEXT_PUBLIC_API_URL, CORS_ORIGINS"],
     ["Seguridad", "ENCRYPTION_KEY, NODE_ENV, PORT"]],
    widths=[3.5, 11.5])
doc.add_heading("11.2 Despliegue", level=2)
placeholder("[ FALTA DEFINIR: donde corre la API en produccion (host, dominio, HTTPS). Frontend en Vercel. ]")
doc.add_heading("11.3 Assets estaticos y base de datos", level=2); placeholder()

# ======================================================================
# CAPITULO 12
# ======================================================================
chapter(12, "Operacion y mantenimiento")
guidance_box(
    objetivo="Cubrir el dia a dia y el soporte del sistema.",
    contenido="Scripts de seed (orden recomendado), crons (FX, expiracion pedidos/pagos, sync lotes), "
              "logs/errores normalizados, respaldos de BD, regeneracion de Prisma, checklist post-deploy.",
    extension="2 paginas",
    tablas="Catalogo de crons (plantilla incluida).")
doc.add_heading("12.1 Seeds y datos iniciales", level=2); placeholder()
doc.add_heading("12.2 Tareas programadas (cron)", level=2)
add_table(
    ["Cron", "Frecuencia", "Responsabilidad", "Archivo"],
    [["Sync tasas FX", "Cada hora", "MXN->USD desde ExchangeRate-API", "tasas-cambio-sync.service.ts"],
     ["Expiracion de pagos", "Diaria", "Marca pagos pendientes vencidos", "pagos.service.ts"],
     ["Expiracion de pedidos", "Programada", "Marca ordenes no pagadas como expiradas", "pedidos.service.ts"],
     ["Sync de lotes", "Programada", "Reconcilia inventario por lote", "lotes.service.ts"]],
    widths=[3, 2.5, 6, 3.5])
doc.add_heading("12.3 Respaldos, logs y checklist post-deploy", level=2)
placeholder("[ FALTA DEFINIR: estrategia de respaldo/restore de Neon y monitoreo. ]")

# ======================================================================
# CAPITULO 13
# ======================================================================
chapter(13, "Limitaciones, deuda tecnica y trabajo futuro")
guidance_box(
    objetivo="Declarar con honestidad el estado real y guiar la evolucion.",
    contenido="Funcionalidades parciales o planeadas, riesgos asociados y recomendaciones priorizadas.",
    extension="2 paginas",
    tablas="Tabla de pendientes (incluida abajo).")
add_table(
    ["Pendiente", "Estado", "Riesgo / Impacto", "Recomendacion"],
    [["FedEx/Easyship directo", "No implementado (via SkydropX)", "Dependencia de un agregador", "Evaluar integracion directa si crece volumen"],
     ["i18n backend", "Stub", "Bajo (frontend cubre)", "Completar o retirar el modulo"],
     ["Factura CFDI (MX)", "Parcial", "Cumplimiento fiscal MX", "Definir PAC/timbrado antes de produccion"],
     ["Resolucion de disputas (UI)", "Parcial", "Operacion de soporte", "Completar flujo de gestion"],
     ["Sales tax US por estado", "No implementado", "Cumplimiento fiscal US", "Definir nexus del negocio"],
     ["Export de reportes", "Parcial", "Bajo", "Agregar export PDF/CSV"]],
    widths=[3.2, 3.2, 4, 4.6])

# ======================================================================
# CAPITULO 14
# ======================================================================
chapter(14, "Anexos")
guidance_box(
    objetivo="Material de referencia para consulta rapida.",
    contenido="Glosario, indice de endpoints (apuntar a Swagger /api/docs), tabla completa de variables "
              "de entorno, diagrama ER ampliado.",
    extension="2 paginas")
doc.add_heading("Anexo A. Glosario", level=2); placeholder()
doc.add_heading("Anexo B. Indice de endpoints (Swagger /api/docs)", level=2); placeholder()
doc.add_heading("Anexo C. Diagrama ER ampliado", level=2); placeholder("[ Insertar ER completo ]")

# ======================================================================
# Guardar
# ======================================================================
out_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "Manual-Tecnico-Marketplace-Residencia.docx")
doc.save(out_path)
print("OK ->", out_path)
